import os, tempfile, zipfile, re
from django.shortcuts import render
from django.shortcuts import render_to_response
from .models import Post
from django.utils.encoding import smart_str
from django.core.servers.basehttp import FileWrapper
from django.http import StreamingHttpResponse
from django.http import FileResponse
from django.conf import settings
import mimetypes
from docx import *
from time import localtime
from shutil import copyfile

# Create your views here.
from django.http import HttpResponse

# Запрос страницы с ссылками
def page_link(request):
	posts = Post.objects.order_by('name_document')
	return render(request, 'saledoc/page_link.html', {'posts': posts})

def data_vremya():
	# From PC date read.
	datatime = localtime()

	data1 = datatime[0]
	data2 = datatime[1]
	data3 = datatime[2]
	data4 = datatime[3]
	data5 = datatime[4]
	data = []
	data.append(data1)
	data.append(data2)
	data.append(data3)
	data.append(data4)
	data.append(data5)

	datavalue = ''.join(str(v) for v in data)

	return datavalue

def create_copy_fileshablon(namedoc):
	copyfile('saledoc/test{}.docx'.format(int(namedoc)), 'saledoc/{}.docx'.format(data_vremya()))
	path_to_newfile = 'saledoc/{}.docx'.format(data_vremya())
	return path_to_newfile

def documents(request, arg2, arg3):
	if(request.GET.get('mybtn')):
		document_select = "7"
		namedoc = arg3
		data_dogovora = request.GET.get('data_dogovor')
		CustOrgShortName = request.GET.get('CustOrgShortName')
		CustBankDetails = request.GET.get('CustBankDetails')
		CustOrgFullName = request.GET.get('CustOrgFullName')
		RCustFullName = request.GET.get('RCustFullName')
		CustInitials = request.GET.get('CustInitials')
		place = request.GET.get('place')
		place_dogovor_writing = request.GET.get('place_dogovor_writing')
		CountOrgShortName = request.GET.get('CountOrgShortName')
		VatRate = request.GET.get('VatRate')
		AgrCurrency = request.GET.get('AgrCurrency')
		CurrencyPayment = request.GET.get('CurrencyPayment')
		podpis_name = request.GET.get('podpis_name')
		CustPost = request.GET.get('CustPost')
		RCustPost = request.GET.get('RCustPost')
		CustLastName = request.GET.get('CustLastName')
		CustBase = request.GET.get('CustBase')
		number_dogovor = request.GET.get('number_dogovor')
		CostNum = request.GET.get('CostNum')
		VatRate_keyboard = request.GET.get('VatRate_keyboard')
		
		AgrLaw = request.GET.get('AgrLaw')
		ProductName = request.GET.get('ProductName')
		ProductName_type = request.GET.get('ProductName_type')
		CostPrePayment = request.GET.get('CostPrePayment')
		CostPrePayment_type = request.GET.get('CostPrePayment_type')
		ProductNameDetail = request.GET.get('ProductNameDetail')
		
		new_file_is = create_copy_fileshablon(namedoc)

		writedata_indocfile(new_file_is, data_dogovora, RCustFullName, CustOrgShortName, CustBankDetails, CustOrgFullName, CustInitials, place, CountOrgShortName, VatRate, AgrCurrency, CurrencyPayment, podpis_name, CustPost, RCustPost, CustLastName, CustBase, number_dogovor, CostNum, VatRate_keyboard, place_dogovor_writing, AgrLaw, ProductName, ProductName_type, CostPrePayment, CostPrePayment_type, ProductNameDetail)
	else:
		new_file_is = str()	
		document_select = arg3
	return render_to_response('saledoc/page_content.html', {'document_selected': document_select, 'news_file' : new_file_is})

def request_page(request):
	datas = request.GET.get('CountOrgShortName')
	print(datas)
	return render(request, 'saledoc/page_download.html', {'inform': datas})

def send_file(request, fileURL):
	# send file "fileURL" as response
	filename = 'saledoc/{}'.format(fileURL)
	download_name =  os.path.basename(filename)
	chunk_size = 8192
	response = StreamingHttpResponse(FileWrapper(open(filename	, 'rb'), chunk_size), content_type=mimetypes.guess_type(filename)[0])
	response['Content-Length'] = os.path.getsize(filename)
	response['Content-Disposition'] = "attachment; filename=%s" % download_name
	return response

def string_to_writenumber(summa_dogovor):
	# change format style from number into write number.
	ZERO = "ноль"
	
	L_1_HE = HE = [ZERO, "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять", "десять",
	    "одинадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
	    "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
	    
	L_1_SHE = SHE = [ZERO, "одна", "две"] + L_1_HE[3:]
	    
	L_10 = [ZERO, "десять", "двадцать", "тридцать", "сорок", "пятьдесят",
	    "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
	    
	L_100 = [ZERO, "сто", "двести", "триста", "четыреста", "пятьсот",
	    "шестьсот", "семьсот", "восемьсот", "девятьсот"]
	    
	N_ROUBLE = "рубл"
	N_COP = "копе"
	RUR = (N_ROUBLE, N_COP)
	
	N_DOLLAR = "доллар"
	N_CENT = "цент"
	USD = (N_DOLLAR, N_CENT)
	
	GENDER = {
	    N_ROUBLE: HE,
	    N_COP: SHE,
	    N_DOLLAR: HE,
	    N_CENT: HE,
	} 
	
	N_1000 = "тысяч"
	N_MILLION = "миллион"
	N_BILLION = "миллиард"
	
	ENDINGS = {
	    N_ROUBLE:     ["ей", "ь", "я", "я", "я"] + 30 * ["ей"],
	    N_COP:    ["ек", "йка", "йки", "йки", "йки"] + 30 * ["ек"],
	    N_DOLLAR:    ["ов", "", "а", "а", "а"] + 30 * ["ов"],
	    N_CENT:    ["ов", "", "а", "а", "а"] + 30 * ["ов"],
	    N_1000:    ["", "а", "и", "и", "и"] + 30 * [""],
	    N_MILLION:    ["ов", "", "а", "а", "а"] + 30 * ["ов"],
	    N_BILLION:    ["ов", "", "а", "а", "а"] + 30 * ["ов"],
	}
	
	def write_1_to_999(n, gender_digits=HE):
	    assert n <= 999
	    
	    if n == 0:
	        return ZERO
	    
	    n_100 = n // 100
	    n_10 = n % 100 // 10 
	    n_1 = n % 10
	    
	    res = []
	    res.append(L_100[n_100])
	    
	    if n_10 == 1:
	        res.append(gender_digits[10 * n_10 + n_1])
	    else:
	        res.append(L_10[n_10])
	        res.append(gender_digits[n_1])
	    return " ".join([s for s in res if s != ZERO])
	
	def ending_index(n):
	    n_2 = n % 100
	    return n_2 if n_2 < 20 else n_2 % 10
	
	def form_group_name(group, n):
	    return group + ENDINGS[group][ending_index(n)]
	    
	def form_group(group, n, gender_digits=HE):
	    return ("%s %s" % (write_1_to_999(n, gender_digits), form_group_name(group, n))) if n else ZERO
	
	def write_number(n, gender_digits=HE):
	    assert type(n) in (int, float)
	    if n == 0:
	        return ZERO
	    
	    n_3 = n % 10 ** 3
	    n_6 = n % 10 ** 6 // 10 ** 3
	    n_9 = n % 10 ** 9 // 10 ** 6
	    n_12 = n % 10 ** 12 // 10 ** 9
	    
	    # print n_12, n_9, n_6, n_3
	    res = []
	    
	    res.append(form_group(N_BILLION, n_12))
	    res.append(form_group(N_MILLION, n_9))
	    res.append(form_group(N_1000, n_6, SHE))
	    res.append(write_1_to_999(n_3, gender_digits))
	    
	    return ", ".join([s for s in res if s != ZERO])
	'''
	def write_price(n_rub, n_cop=0, currency=RUR):
	    rub_id, cop_id = currency
	    n = int(n_rub)
	    res = []
	    res.append("%s %s" % (write_number(n, GENDER[rub_id]), form_group_name(rub_id, n)))
	    res.append(form_group(cop_id, n_cop, GENDER[cop_id]))
	    
	    return " и ".join([s for s in res if s != ZERO])
	'''  
	summapropisu = write_number(summa_dogovor)
	return summapropisu

def replace_string(filename, old_info, new_info):
    doc = Document(filename)
    for p in doc.paragraphs:
        if old_info in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if old_info in inline[i].text:
                    text = inline[i].text.replace(old_info, new_info)
                    inline[i].text = text
            print (p.text)

    doc.save(filename)
    return 1
   
def writedata_indocfile(new_file_is, data_dogovora, CustOrgFullName, CustOrgShortName, CustBankDetails, RCustFullName, CustInitials, place, CountOrgShortName, VatRate, AgrCurrency, CurrencyPayment, podpis_name, CustPost, RCustPost, CustLastName, CustBase, number_dogovor, CostNum, VatRate_keyboard, place_dogovor_writing, AgrLaw, ProductName, ProductName_type, CostPrePayment, CostPrePayment_type, ProductNameDetail):
	
	# Generating number of document
	replace_string(new_file_is, 'AgrNum_', 'ДОГОВОР №{}'.format(data_vremya()))

	# Check param "CountOrgShortName" and replace in doc file "CountOrgShortName_"
	if int(CountOrgShortName) == 1:
		replace_string(new_file_is, 'CountOrgShortName_', 'ЧУП Промвад') 
	elif int(CountOrgShortName) == 2:
		replace_string(new_file_is, 'CountOrgShortName_', 'ООО Промвад Софт')
	else:
		replace_string(new_file_is, 'CountOrgShortName_', 'ПРВ Инжиниринг')

	# Check param "podpis_name" and replace in doc file "CountPost_,RCountPost_,CountLastName_, CountInitials_, RCountFullName_, RCountFullName_, CountBase_, CountBaseNum_, CountBaseDate_, CountOrgFullName_, CountBankDetails_"
	if int(podpis_name) == 1:
		replace_string(new_file_is, 'CountPost_', 'Директор') 
		replace_string(new_file_is, 'RCountPost_', 'директора')
		replace_string(new_file_is, 'CountLastName_', 'Якубович')
		replace_string(new_file_is, 'CountInitials_', 'Д.М.')
		replace_string(new_file_is, 'RCountFullName_', 'Якубовича Дениса Михайловича')
		replace_string(new_file_is, 'CountBase_', 'Устава')
		replace_string(new_file_is, 'CountBaseNum_', '')
		replace_string(new_file_is, 'CountBaseDate_', '')
		replace_string(new_file_is, 'CountOrgFullName_', 'Частное унитарное предприятие Промвад')
		replace_string(new_file_is, 'CountBankDetails_', '220073, г. Минск, ул. Ольшевского, 20/11 Банк: ЗАО «Идея Банк» г.Минск,  БИК 153001755 УНП 191448150, ОКПО 379817905')
	elif int(podpis_name) == 2:
		replace_string(new_file_is, 'CountPost_', 'Заместитель директора по производству') 
		replace_string(new_file_is, 'RCountPost_', 'заместителя директора по производству')
		replace_string(new_file_is, 'CountLastName_', 'Ковалев')
		replace_string(new_file_is, 'CountInitials_', 'С.Н.')
		replace_string(new_file_is, 'RCountFullName_', 'Ковалева Сергея Николаевича')
		replace_string(new_file_is, 'CountBase_', 'доверенности')
		replace_string(new_file_is, 'CountBaseNum_', '01-27/4336')
		replace_string(new_file_is, 'CountBaseDate_', '25.06.2012')
		replace_string(new_file_is, 'CountOrgFullName_', 'Частное унитарное предприятие Промвад')
		replace_string(new_file_is, 'CountBankDetails_', '220073, г. Минск, ул. Ольшевского, 20/11 Банк: ЗАО «Идея Банк» г.Минск,  БИК 153001755 УНП 191448150, ОКПО 379817905')
	elif int(podpis_name) == 3:
		replace_string(new_file_is, 'CountPost_', 'Директор') 
		replace_string(new_file_is, 'RCountPost_', 'директора')
		replace_string(new_file_is, 'CountLastName_', 'Мозолевский')
		replace_string(new_file_is, 'CountInitials_', 'В.В.')
		replace_string(new_file_is, 'RCountFullName_', 'Мозолевского Виталия Владимировича')
		replace_string(new_file_is, 'CountBase_', 'Устава')
		replace_string(new_file_is, 'CountBaseNum_', '')
		replace_string(new_file_is, 'CountBaseDate_', '')
		replace_string(new_file_is, 'CountOrgFullName_', 'Общество с ограниченной ответственностью «Промвад Софт»')
		replace_string(new_file_is, 'CountBankDetails_', '220073, г. Минск, ул. Ольшевского, 20/11 Банк: ЗАО «Идея Банк» г.Минск,  БИК 153001755 УНП 191448150, ОКПО 379817905')
	elif int(podpis_name) == 4:
		replace_string(new_file_is, 'CountPost_', 'Заместитель директора по коммерческим вопросам') 
		replace_string(new_file_is, 'RCountPost_', 'заместителя директора по коммерческим вопросам')
		replace_string(new_file_is, 'CountLastName_', 'Лиштван')
		replace_string(new_file_is, 'CountInitials_', 'Ю.Н.')
		replace_string(new_file_is, 'RCountFullName_', 'Лиштвана Юрия Николаевича')
		replace_string(new_file_is, 'CountBase_', 'доверенности')
		replace_string(new_file_is, 'CountBaseNum_', '1')
		replace_string(new_file_is, 'CountBaseDate_', '03.01.2014')
		replace_string(new_file_is, 'CountOrgFullName_', 'Общество с ограниченной ответственностью «Промвад Софт»')
		replace_string(new_file_is, 'CountBankDetails_', '220073, г. Минск, ул. Ольшевского, 20/11 Банк: ЗАО «Идея Банк» г.Минск,  БИК 153001755 УНП 191448150, ОКПО 379817905')
	elif int(podpis_name) == 5:
		replace_string(new_file_is, 'CountPost_', 'Директор') 
		replace_string(new_file_is, 'RCountPost_', 'директора')
		replace_string(new_file_is, 'CountLastName_', 'Кутень')
		replace_string(new_file_is, 'CountInitials_', 'И.С.')
		replace_string(new_file_is, 'RCountFullName_', 'Кутеня Ивана Семеновича')
		replace_string(new_file_is, 'CountBase_', 'Устава')
		replace_string(new_file_is, 'CountBaseNum_', '')
		replace_string(new_file_is, 'CountBaseDate_', '')
		replace_string(new_file_is, 'CountOrgFullName_', 'ПРВ Инжиниринг')
		replace_string(new_file_is, 'CountBankDetails_', '220073, г. Минск, ул. Ольшевского, 20/11 Банк: ЗАО «Идея Банк» г.Минск,  БИК 153001755 УНП 191448150, ОКПО 379817905')
	else:
		replace_string(new_file_is, 'CountPost_', 'Заместитель директора по производству ') 
		replace_string(new_file_is, 'RCountPost_', 'заместителя директора по производству ')
		replace_string(new_file_is, 'CountLastName_', 'Ковалев')
		replace_string(new_file_is, 'CountInitials_', 'С.Н.')
		replace_string(new_file_is, 'RCountFullName_', 'Ковалева Сергея Николаевича ')
		replace_string(new_file_is, 'CountBase_', 'доверенности')
		replace_string(new_file_is, 'CountBaseNum_', '1')
		replace_string(new_file_is, 'CountBaseDate_', '31.03.2015')
		replace_string(new_file_is, 'CountOrgFullName_', 'ПРВ Инжиниринг')
		replace_string(new_file_is, 'CountBankDetails_', '220073, г. Минск, ул. Ольшевского, 20/11 Банк: ЗАО «Идея Банк» г.Минск,  БИК 153001755 УНП 191448150, ОКПО 379817905')

	# Check param "place" and replace in doc file "AgrCity_"
	document.add_heading('4. Место составления договора', level=1)
	if int(place) == 1: 
		replace_string(new_file_is, 'AgrCity_', 'Минск')
	elif int(place) == 2:
		replace_string(new_file_is, 'AgrCity_', 'Москва')
	else:
		replace_string(new_file_is, 'AgrCity_', place_dogovor_writing)

	# Generating data of document
	replace_string(new_file_is, 'AgrDate_', '{}'.format(data_vremya()))

	# Check param "CustOrgFullName, CustOrgShortName, CustBankDetails" and replace in doc file "CustOrgFullName_, CustOrgShortName_, CustBankDetails_"
	replace_string(new_file_is, 'CustOrgFullName_', CustOrgFullName)
	replace_string(new_file_is, 'CustOrgShortName_', CustOrgShortName)
	replace_string(new_file_is, 'CustBankDetails_', CustBankDetails)

	# Check param "CustPost, RCustPost, CustLastName, RCustFullName, CustInitials" and replace in doc file "CustPost_, RCustPost_, CustLastName_, RCustFullName_, CustInitials_"
	# Proverka chemu raven parametr CustBase i formirovanie informacii o predstavitele zakazchika.
	replace_string(new_file_is, 'CustPost_', CustPost)
	replace_string(new_file_is, 'RCustPost_', RCustPost)
	replace_string(new_file_is, 'CustLastName_', CustLastName)
	replace_string(new_file_is, 'RCustFullName_', RCustFullName)
	replace_string(new_file_is, 'CustInitials_', CustInitials)

	if int(CustBase) == 1: 
		replace_string(new_file_is, 'CustBase_', 'Устава')
	else:
		replace_string(new_file_is, 'CustBase_', 'доверенности №')

	# Check param "VatRate" and replace in doc file "VatRate_"
	if int(VatRate) == 1:
		rate = 0 
		replace_string(new_file_is, 'VatRate_', 'Стоимость работ без НДС')
	elif int(VatRate) == 2:
		rate = 18
		replace_string(new_file_is, 'VatRate_', '18% (для РФ)')
	elif int(VatRate) == 3:
		rate = 20
		replace_string(new_file_is, 'VatRate_', '20% (для РБ')
	else:
		replace_string(new_file_is, 'VatRate_', VatRate_keyboard)
		rate = int(VatRate_keyboard)

	# Check param "AgrCurrency" and replace in doc file "AgrCurrency_".
	if int(AgrCurrency) == 1:
		replace_string(new_file_is, 'AgrCurrency_', 'долларах США') 
		replace_string(new_file_is, 'CountRS_', 'р/с 3012082311015 SWIFT: SOMA BY 22 Банк-корреспондент: Deutsche bank AG, Frankfurt am Main SWIFT DEUTDEFF Номер счета, валюта: 10094779510000')
	elif int(AgrCurrency) == 2:
		replace_string(new_file_is, 'AgrCurrency_', 'Евро')
		replace_string(new_file_is, 'CountRS_', 'р/с 3012082311015 SWIFT: SOMA BY 22 Банк-корреспондент: Deutsche bank AG, Frankfurt am Main SWIFT DEUTDEFF Номер счета, валюта: 10094779510000')
	elif int(AgrCurrency) == 3:
		replace_string(new_file_is, 'AgrCurrency_', 'российских рублях')
		replace_string(new_file_is, 'CountRS_', 'р/с 3012082311015 SWIFT: SOMA BY 22 Банк-корреспондент: Deutsche bank AG, Frankfurt am Main SWIFT DEUTDEFF Номер счета, валюта: 10094779510000')
	else:
		replace_string(new_file_is, 'AgrCurrency_', 'белорусских рублях')
		replace_string(new_file_is, 'CountRS_', 'р/с 3012082311015 SWIFT: SOMA BY 22 Банк-корреспондент: Deutsche bank AG, Frankfurt am Main SWIFT DEUTDEFF Номер счета, валюта: 10094779510000')

	if int(CurrencyPayment) == 1: 
		replace_string(new_file_is, 'CurrencyPayment_', 'доллары США')
	elif int(CurrencyPayment) == 2:
		replace_string(new_file_is, 'CurrencyPayment_', 'Евро')
	elif int(CurrencyPayment) == 3:
		replace_string(new_file_is, 'CurrencyPayment_', 'российский рубль')
	else:
		replace_string(new_file_is, 'CurrencyPayment_', 'белорусский рубль')

		
	if int(CurrencyPayment) == 3 and int(AgrCurrency) == 3:
		replace_string(new_file_is, 'CurrencyText_', 'В случае изменения курса российского рубля по отношению к доллару США на день выставления счета по сравнению с курсом на дату заключения настоящего договора,  Исполнитель вправе проиндексировать сумму настоящего договора в части невыставленных счетов путем пересчета суммы счета в доллары США по курсу ЦБРФ на день подписания  настоящего договора и последующего пересчета в российские рубли по курсу ЦБРФ на день выставления счета. Кроме того, в случае нарушения срока оплаты согласно настоящего договора, Исполнитель вправе пересчитать сумму выставленного, но неоплаченного счета и перевыставить счет с учетом корректировки.')
	elif int(CurrencyPayment) == 4 and int(AgrCurrency) == 4:
		replace_string(new_file_is, 'CurrencyText_', 'В случае изменения курса белорусского рубля по отношению к доллару США на день выставления счета по сравнению с курсом на дату заключения настоящего договора,  Исполнитель вправе проиндексировать сумму настоящего договора в части невыставленных счетов путем пересчета суммы счета в доллары США по курсу НБРБ на день подписания  настоящего договора и последующего пересчета в белорусские рубли по курсу НБРБ на день выставления счета. Кроме того, в случае нарушения срока оплаты согласно настоящего договора, Исполнитель вправе пересчитать сумму выставленного, но неоплаченного счета и перевыставить счет с учетом корректировки.')
	else:
		replace_string(new_file_is, 'CurrencyText_', '')
			
	# Check param "CostNum" and replace in doc file "CostNum_".
	document_Summa = string_to_writenumber(int(CostNum))
	replace_string(new_file_is, 'CostNum_', document_Summa)
	NDSNum = int(CostNum) * int(rate)
	NDSNum_write = string_to_writenumber(int(NDSNum))
	replace_string(new_file_is, 'NDSNum_', NDSNum_write)
	
	# Check param "AgrLaw" and replace in doc file "AgrLaw_".
	if int(AgrLaw) == 1: 
		replace_string(new_file_is, 'AgrLaw_', 'Арбитражном суде г. Москвы')
	else:
		replace_string(new_file_is, 'AgrLaw_', 'Хозяйственном суде по месту нахождения Истца')

	# Check param "ProductName" and replace in doc file "ProductName_".	
	if int(ProductName) == 1: 
		replace_string(new_file_is, 'ProductName_', 'разработку научно технической продукции')
	elif int(ProductName) == 2:
		replace_string(new_file_is, 'ProductName_', 'проведение опытно конструкторских работ')
	elif int(ProductName) == 3:
		replace_string(new_file_is, 'ProductName_', 'разработку программного обеспечения')
	else:
		replace_string(new_file_is, 'ProductName_', ProductName_type)
	
	# Check param "ProductName" and replace in doc file "ProductName_".	
	if int(CostPrePayment) == 1: 
		replace_string(new_file_is, 'CostPrePayment_', '70 (семидесяти)')
	elif int(CostPrePayment) == 2:
		replace_string(new_file_is, 'CostPrePayment_', '50 (пятидесяти)')
	else:
		replace_string(new_file_is, 'CostPrePayment_', CostPrePayment_type)
		
	replace_string(new_file_is, 'ProductNameDetail_', ProductNameDetail)
			
	return



